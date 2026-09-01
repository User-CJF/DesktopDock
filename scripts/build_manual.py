from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "DesktopDock_完整项目说明书.docx"
SCREENSHOT_DESKTOP = next((item for item in [ROOT / "dist" / "electron-widget-dock.png", ROOT / "release" / "electron-widget-dock.png"] if item.exists()), ROOT / "dist" / "electron-widget-dock.png")
SCREENSHOT_SETTINGS = next((item for item in [ROOT / "dist" / "electron-widget-settings.png", ROOT / "release" / "electron-widget-settings.png"] if item.exists()), ROOT / "dist" / "electron-widget-settings.png")
VERSION = "0.2.0"
EXE_NAME = f"DesktopDock-{VERSION}-Portable.exe"
EXE_SHA256 = "CF7887A71183D5540D10F17421552A66EDF9FEB57982746C0289CFF00FF7604F"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, value, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(value))
    run.bold = bold
    run.font.name = "Microsoft YaHei UI"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, value in enumerate(headers):
        set_cell_text(header.cells[index], value, bold=True, color="FFFFFF")
        set_cell_shading(header.cells[index], "1677FF")
        if widths:
            header.cells[index].width = Cm(widths[index])
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            if widths:
                cells[index].width = Cm(widths[index])
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_paragraph(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.35
    if bold_prefix and text.startswith(bold_prefix):
        paragraph.add_run(bold_prefix).bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.add_run(item)


def add_steps(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    paragraph.add_run(" 页")


doc = Document()
doc.core_properties.title = "桌面舱 DesktopDock 完整项目说明书"
doc.core_properties.subject = "DesktopDock 0.2.0 产品、功能、架构、测试与运维"
doc.core_properties.author = "DesktopDock Contributors"
doc.core_properties.comments = "本文件是当前版本唯一项目说明书"

section = doc.sections[0]
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.25)
section.right_margin = Cm(2.25)
section.header_distance = Cm(0.8)
section.footer_distance = Cm(0.8)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei UI"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor(40, 50, 62)
for name, size, color in [("Title", 34, "17202A"), ("Subtitle", 15, "5D6A78"), ("Heading 1", 19, "0A5FAE"), ("Heading 2", 14, "17202A"), ("Heading 3", 11.5, "30506F")]:
    style = styles[name]
    style.font.name = "Microsoft YaHei UI"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)

header = section.header.paragraphs[0]
header.text = "桌面舱 DesktopDock  ·  完整项目说明书  ·  v0.2.0"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header.runs[0].font.size = Pt(8)
header.runs[0].font.color.rgb = RGBColor(112, 126, 141)
add_page_number(section.footer.paragraphs[0])

# Cover
cover_mark = doc.add_paragraph()
cover_mark.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_mark.paragraph_format.space_before = Pt(48)
run = cover_mark.add_run("▦")
run.font.name = "Segoe Fluent Icons"
run.font.size = Pt(42)
run.font.color.rgb = RGBColor(22, 119, 255)
title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("桌面舱 DesktopDock")
subtitle = doc.add_paragraph(style="Subtitle")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("完整项目说明书")
tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline.paragraph_format.space_before = Pt(8)
tagline.add_run("桌面仓 · 用户分类 · 文件格子 · 待办 · 天气 · 媒体 · Windows 桌面宿主层").font.color.rgb = RGBColor(83, 101, 120)
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(90)
meta.add_run(f"版本 {VERSION}\n生成日期 2026-09-01\n唯一有效项目文档").font.size = Pt(10)
doc.add_page_break()

add_heading(doc, "文档说明", 1)
add_paragraph(doc, "本说明书是 DesktopDock 当前版本的唯一完整项目文档，整合并替代早期三份 UI 开发说明书和一份旧完整开发说明书。文档以已经构建并通过独立冒烟测试的 0.2.0 便携版为准；附图、功能清单、权限边界和限制均对应实际实现，不把设计设想写成已交付功能。")
add_table(doc, ["项目", "内容"], [
    ("交付程序", EXE_NAME),
    ("程序 SHA-256", EXE_SHA256),
    ("远程仓库", "https://github.com/User-CJF/DesktopDock.git"),
    ("目标系统", "Windows 10/11 x64"),
    ("数据策略", "本地优先；无账号；不上传本机路径、快捷方式或待办"),
], [4.2, 12.5])

add_heading(doc, "目录", 1)
for line in [
    "1. 项目概述与设计原则", "2. 交付、安装与桌面层行为", "3. 信息架构与视觉规范", "4. 桌面仓与快捷方式分类", "5. 文件、待办、天气与媒体", "6. 设置、数据与安全", "7. 技术架构与核心流程", "8. 构建、测试与验收", "9. 运维、Git 与故障处理", "附录：快捷键、验收表与已知边界",
]:
    add_paragraph(doc, line)
doc.add_page_break()

add_heading(doc, "1　项目概述与设计原则", 1)
add_heading(doc, "1.1 产品定位", 2)
add_paragraph(doc, "DesktopDock 是面向 Windows 10/11 的本地桌面整理与效率工具。它不是传统的全屏“软件管理器”，而是固定在桌面右侧的窄幅状态栏：只有回到桌面时可见，普通应用窗口会自然覆盖它。主流程围绕真实桌面快捷方式展开，所有分类由用户亲自创建。")
add_heading(doc, "1.2 四条产品原则", 2)
add_table(doc, ["原则", "落实方式"], [
    ("原生优先", "使用 Windows Shell 图标、系统媒体会话、系统通知、回收站和桌面宿主层。"),
    ("本地优先", "快捷方式、分类、文件索引、待办、设置和缓存均保存在当前用户本机。"),
    ("克制扩展", "固定桌面仓，其他分类由用户决定；功能模块可单独关闭。"),
    ("轻量常驻", "500px 窄栏、托盘常驻、增量索引、天气缓存，不占用应用工作区。"),
], [3.4, 13.3])
add_heading(doc, "1.3 核心验收目标", 2)
add_bullets(doc, [
    "获取当前用户桌面与公共桌面的真实快捷方式名称和 Windows 图标。",
    "把当前用户桌面普通快捷方式安全移入桌面仓，系统图标不受影响。",
    "用户可创建、改名、删除分类，可拖动快捷方式，也可用菜单替代拖动。",
    "Dock 绑定 Windows 桌面宿主层，不置顶、不贴到其他应用右侧。",
    "提供可工作的文件、待办、天气、媒体和分组设置模块。",
    "交付一个免安装便携 EXE，并支持开机自动启动。",
])

add_heading(doc, "2　交付、安装与桌面层行为", 1)
add_heading(doc, "2.1 最终交付目录", 2)
add_paragraph(doc, "最终本地目录只保留单文件便携程序、本说明书、.git 与 .gitignore。完整源码保存在 GitHub main 分支，便携 EXE 由 .gitignore 排除，不提交到源码历史。")
add_heading(doc, "2.2 免安装运行", 2)
add_steps(doc, [
    f"把 {EXE_NAME} 放到长期保存的位置。",
    "双击运行。第一次启动会初始化本地数据库、图标缓存和快捷方式保管目录。",
    "如 Windows SmartScreen 提示未知发布者，核对本说明书中的 SHA-256；当前构建未配置商业代码签名证书。",
    "要开机启动，进入“设置 → 常规”开启“开机自动启动”。",
    "关闭或隐藏后程序保留在托盘；彻底退出请使用设置页或托盘菜单的“退出”。",
])
add_heading(doc, "2.3 桌面宿主层", 2)
add_paragraph(doc, "Windows 版启动后通过受控 Win32 调用把 Electron 窗口挂接到 Explorer 的 Progman/WorkerW 桌面宿主层。结果是 Dock 只在桌面右侧出现；切换到浏览器、编辑器或游戏时，应用窗口会覆盖 Dock。Dock 不申请 always-on-top，也不注册系统 AppBar，因此不会挤压其他窗口的可用区域。Explorer 重启或显示器参数改变后，程序会重新尝试附着。")
add_paragraph(doc, "在极少数精简版 Shell、第三方桌面或 Explorer 尚未就绪的环境中，附着可能回退到非置顶右侧窗口；托盘和快捷键仍可控制显示。", bold_prefix="在极少数")

add_heading(doc, "3　信息架构与视觉规范", 1)
add_heading(doc, "3.1 页面结构", 2)
add_paragraph(doc, "界面固定宽度约 500 CSS 像素，纵向占满当前显示器工作区。顶部是品牌、时间、刷新和隐藏；其下是当前模块搜索；中部是唯一滚动区；底部五个入口依次为桌面、待办、文件、组件和设置。此结构减少双列小组件同时争抢注意力的问题。")
add_table(doc, ["模块", "主要职责"], [
    ("桌面", "桌面仓、自建分类、真实快捷方式、拖放与归类菜单。"),
    ("待办", "任务详情、颜色、截止、提醒、重复、附件、排序和批量操作。"),
    ("文件", "目录收藏、本地搜索、排序、缩略图和原生文件操作。"),
    ("组件", "天气与 Windows 媒体会话。"),
    ("设置", "常规、外观、文件格子、功能格子。"),
], [3.2, 13.5])
add_heading(doc, "3.2 视觉规范", 2)
add_bullets(doc, [
    "Windows 原生深色为默认主题，同时支持浅色与跟随系统。",
    "主强调色为 Windows 蓝；分类仅用 1px 色线和轻量色块识别，不使用花哨渐变文字。",
    "字体优先 Segoe UI Variable、Segoe UI、Microsoft YaHei UI；图标来自 Segoe Fluent Icons。",
    "触控目标通常为 30–40px；键盘焦点使用 2px 蓝色外框；支持 prefers-reduced-motion。",
    "页面只允许纵向滚动，测试保证 500px 宽度下无横向溢出。",
])
if SCREENSHOT_DESKTOP.exists():
    doc.add_picture(str(SCREENSHOT_DESKTOP), width=Inches(3.55))
    caption = doc.add_paragraph("图 3-1　桌面仓主界面（真实 Windows 图标）")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True

add_heading(doc, "4　桌面仓与快捷方式分类", 1)
add_heading(doc, "4.1 扫描范围与图标", 2)
add_paragraph(doc, "桌面模块扫描当前用户桌面、公共桌面和 DesktopDock 专用快捷方式保管区，只接受 .lnk、.url 与 .appref-ms。主进程解析快捷方式显式图标、目标文件图标或 Shell 后备图标，再转换为 PNG Data URL 返回渲染进程。渲染进程永远不接收真实完整路径。")
add_heading(doc, "4.2 自动收纳", 2)
add_paragraph(doc, "“自动收纳桌面快捷方式”默认开启。启动时，当前用户桌面上的普通快捷方式会移动到当前用户 DesktopDock 数据目录中的 desktop-shortcuts 保管区，桌面仓继续显示并可启动这些快捷方式。此电脑、回收站、控制面板等系统图标由 Explorer 管理，不是普通快捷方式文件，因此不会被移动或删除。")
add_paragraph(doc, "公共桌面快捷方式通常由所有用户共享并可能需要管理员权限，DesktopDock 只读展示，不自动移动。界面会单独显示公共桌面只读数量。")
add_heading(doc, "4.3 用户分类", 2)
add_bullets(doc, [
    "桌面仓是固定入口，不可删除；不再自动生成办公、开发、娱乐等分类。",
    "新分类只要求名称和识别色；所有分类均可改名或删除。",
    "删除分类会把其中快捷方式移回桌面仓，不删除快捷方式本体。",
    "把仓内快捷方式拖到分类卡即可归类；也可从 Windows 桌面直接拖到桌面仓或分类。",
    "每个快捷方式的省略号菜单提供“移动到分类”替代操作，满足不便拖动时的使用需求。",
    "双击快捷方式或选中后按 Enter 可启动目标。",
])
add_heading(doc, "4.4 拖入与冲突策略", 2)
add_paragraph(doc, "从当前用户桌面拖入时使用安全移动，从其他目录拖入时复制快捷方式，避免意外破坏原文件。目标同名时自动生成“名称 (1)”等安全文件名，不覆盖已有内容。一次最多处理 100 个路径，非快捷方式会被跳过并报告数量。")

add_heading(doc, "5　文件、待办、天气与媒体", 1)
add_heading(doc, "5.1 文件格子", 2)
add_bullets(doc, [
    "预置桌面、下载、文档、图片、视频；可收藏和移除其他目录。",
    "本地 SQLite 索引限制扫描深度和单根目录数量，忽略 node_modules、.git、回收站等目录。",
    "支持当前模块搜索，按最近修改、名称或大小排序；支持图标和列表布局。",
    "PNG、JPEG、GIF、BMP、WebP 使用真实缩略图，其他文件显示扩展名标记。",
    "支持拖入复制、打开、在资源管理器中显示、重命名和移到 Windows 回收站。",
    "拖入文件会复制到选中目录，原文件保留；删除调用回收站而非直接永久删除。",
])
add_heading(doc, "5.2 待办", 2)
add_bullets(doc, [
    "任务字段包括标题、备注、八种颜色、截止时间、提醒时间、重复规则和最多 20 个附件路径。",
    "任务可拖动排序；可单项完成/恢复，也可进入批量模式完成或删除所选任务。",
    "提醒到期后由 Electron 调用 Windows 系统通知；提醒仅在 DesktopDock 运行期间触发。",
    "每天、每周或每月重复的任务在首次完成时自动创建下一期，保留备注、颜色和附件。",
    "附件只保存本机路径，不复制或上传文件；路径失效不会删除任务。",
])
add_heading(doc, "5.3 天气", 2)
add_paragraph(doc, "天气由 Open-Meteo 地理编码与预报接口提供，缓存 30 分钟。支持设置页手动城市，也支持天气页请求当前位置；系统定位需要 Windows 位置权限。显示当前温度、体感、湿度、风速、气压、降水概率、未来六小时和七天预报，以及日出、日落、紫外线等已缓存字段。标准、紧凑、逐小时和周预报四种布局可循环切换，背景随晴、多云、降水轻量变化。")
add_heading(doc, "5.4 Windows 媒体", 2)
add_paragraph(doc, "媒体模块通过 Windows Global System Media Transport Controls 读取当前活动会话的曲名、艺术家、播放状态、位置和时长，并通过系统媒体键发送上一首、播放/暂停、下一首、音量降低和音量提高。是否返回封面、随机或循环能力由具体播放器实现；为保持播放器兼容性，当前版本不伪造通用随机/循环命令，也不缓存受版权保护的封面。")

add_heading(doc, "6　设置、数据与安全", 1)
add_heading(doc, "6.1 设置分组", 2)
add_table(doc, ["分组", "设置"], [
    ("常规", "开机启动、自动收纳桌面快捷方式。"),
    ("外观", "深色、浅色、跟随系统。"),
    ("文件格子", "默认布局、模块开关。"),
    ("功能格子", "待办/天气/媒体开关、天气城市与天气布局。"),
], [3.4, 13.3])
if SCREENSHOT_SETTINGS.exists():
    doc.add_picture(str(SCREENSHOT_SETTINGS), width=Inches(3.55))
    caption = doc.add_paragraph("图 6-1　按场景分组后的设置页")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True
add_heading(doc, "6.2 本地数据", 2)
add_table(doc, ["数据", "位置/机制", "可恢复性"], [
    ("分类、待办、设置、索引", "当前用户 AppData 下 data.db", "可导出部分设置；删除 EXE 不会自动删除数据"),
    ("快捷方式", "当前用户 AppData 下 desktop-shortcuts", "可由保管清单恢复；分类删除不删快捷方式"),
    ("天气", "weather-cache.json，30 分钟缓存", "可安全删除，联网后重建"),
    ("图标", "版本化 icon-cache", "可安全删除，Shell 重新提取"),
], [3.2, 7.8, 5.8])
add_heading(doc, "6.3 安全边界", 2)
add_bullets(doc, [
    "Electron 启用 contextIsolation 与 sandbox，关闭 nodeIntegration。",
    "渲染进程通过预加载白名单 IPC 调用主进程；应用、文件和快捷方式使用格式受限的不透明 ID。",
    "重命名拒绝路径分隔符和 Windows 非法字符；删除文件使用 shell.trashItem。",
    "桌面收纳只自动移动当前用户快捷方式，公共桌面只读。",
    "天气是唯一主动联网功能；分类、待办、文件索引和媒体控制不需要云端。",
])

add_heading(doc, "7　技术架构与核心流程", 1)
add_heading(doc, "7.1 技术栈", 2)
add_table(doc, ["层", "技术", "职责"], [
    ("桌面容器", "Electron 43", "窗口、托盘、桌面宿主、通知、Shell 与 IPC"),
    ("界面", "Vue 3 + Vite 8 + 原生 JS/CSS", "稳定壳层、页面渲染、拖放与交互"),
    ("数据", "Node SQLite + JSON 缓存", "分类、待办、索引、设置、缓存"),
    ("系统集成", "Win32 / PowerShell / Windows Shell", "WorkerW、媒体会话、图标、开机启动"),
    ("打包", "electron-builder portable", "单文件 Windows x64 EXE"),
], [2.8, 5.2, 8.8])
add_heading(doc, "7.2 进程职责", 2)
add_bullets(doc, [
    "主进程：所有文件系统、数据库、网络与 Windows 系统调用。",
    "预加载层：只暴露命名清晰、参数受约束的方法。",
    "渲染进程：界面与状态；不能直接读取任意本机路径。",
    "服务层：shortcut-board、desktop-organizer、todo、file-index、weather、media、settings、icon-cache。",
])
add_heading(doc, "7.3 核心流程", 2)
add_table(doc, ["流程", "关键步骤"], [
    ("启动", "单实例锁 → 初始化数据库 → 自动收纳 → 建立索引 → 创建桌面子窗口 → 托盘/快捷键/提醒。"),
    ("拖入分类", "提取安全路径 → 校验扩展名 → 移动或复制到保管区 → 写清单 → 保存分类关联 → 刷新图标。"),
    ("文件删除", "不透明 ID → 主进程解析已索引路径 → Windows 回收站 → 重建索引。"),
    ("提醒", "每 30 秒读取到期待办 → 原子标记已通知 → Windows Notification。"),
    ("媒体", "读取 GSMTC 会话元数据；控制操作映射为 Windows 媒体键。"),
], [3.0, 13.8])

add_heading(doc, "8　构建、测试与验收", 1)
add_heading(doc, "8.1 开发命令", 2)
add_table(doc, ["命令", "用途"], [
    ("pnpm install", "恢复锁定依赖"), ("pnpm run dev", "开发模式"), ("pnpm run check", "语法、12 项测试与生产构建"),
    ("pnpm run smoke", "真实 Electron 500px 窗口冒烟与截图"), ("pnpm run portable:win", "生成单文件便携 EXE"),
], [5.0, 11.8])
add_heading(doc, "8.2 已完成验证", 2)
add_bullets(doc, [
    "12 项 Node 自动测试全部通过：应用索引、分类、桌面整理、快捷方式、文件索引、图标缓存、媒体动作、设置、桌面仓、待办、天气。",
    "Electron 源码冒烟测试通过：桌面仓、文件、待办、天气、媒体、设置、分类表单、IPC 边界和无横向溢出。",
    "0.2.0 单文件 EXE 独立启动冒烟测试返回 passed=true。",
    "视觉截图检查通过：真实图标清晰、导航稳定、设置分组完整、500px 窗口无横向溢出。",
    f"EXE：{EXE_NAME}；大小 98,124,064 bytes；SHA-256：{EXE_SHA256}。",
])
add_heading(doc, "8.3 用户验收清单", 2)
add_table(doc, ["检查项", "预期结果"], [
    ("桌面可见性", "桌面右侧可见；其他应用覆盖；不贴到应用右侧。"),
    ("首次收纳", "个人桌面快捷方式进入仓；此电脑/回收站等系统图标保留。"),
    ("分类", "可创建、拖入、菜单移动、改名、删除；删除后回仓。"),
    ("文件", "搜索/排序/布局/缩略图/拖入/重命名/回收站均正常。"),
    ("待办", "提醒弹窗、重复续建、附件、排序和批量操作正常。"),
    ("天气媒体", "城市或定位天气正常；当前播放器可控制并尽量显示元数据。"),
    ("便携开机", "EXE 无需安装；开启设置后重启 Windows 可自动运行。"),
], [4.0, 12.8])

add_heading(doc, "9　运维、Git 与故障处理", 1)
add_heading(doc, "9.1 Git 规则", 2)
add_paragraph(doc, "远程仓库为 https://github.com/User-CJF/DesktopDock.git，主分支 main。每次修改完成后运行检查、源码冒烟、便携打包与便携冒烟，再提交并推送。根目录 EXE、node_modules、dist、release 和运行数据不进入 Git；GitHub 保存完整源码，本地最终目录使用 sparse-checkout 只保留交付物。")
add_heading(doc, "9.2 常见问题", 2)
add_table(doc, ["现象", "处理"], [
    ("公共桌面快捷方式仍在桌面", "这是权限保护；公共桌面由所有用户共享，DesktopDock 只读展示。"),
    ("定位天气失败", "在 Windows 设置中开启位置服务，或在设置页手动输入城市。"),
    ("媒体没有曲名", "播放器未接入 Windows 媒体会话；控制键仍可能有效，尝试先播放一次。"),
    ("Dock 未显示", "检查托盘并单击“显示桌面舱”；如 Explorer 重启，等待附着或重启 DesktopDock。"),
    ("快捷方式图标为空", "刷新桌面；损坏的 .lnk 会使用后备图标，检查其目标是否存在。"),
    ("开机启动失效", "EXE 移动后重新关闭并开启开机启动，让 Windows 更新路径。"),
], [4.4, 12.4])

add_heading(doc, "附录 A　快捷键", 1)
add_table(doc, ["按键", "作用"], [
    ("Alt + Space", "聚焦 DesktopDock 搜索；冲突时使用备用组合。"),
    ("Alt + D", "显示或隐藏 DesktopDock；冲突时使用 Ctrl + Shift + D。"),
    ("Enter", "在快捷方式获得焦点时启动。"),
    ("Esc", "关闭当前弹层。"),
], [4.2, 12.6])
add_heading(doc, "附录 B　当前边界", 1)
add_bullets(doc, [
    "Windows 系统图标由 Explorer 管理，不会进入桌面仓，也不会被 DesktopDock 删除。",
    "公共桌面快捷方式仅显示，不自动移除；管理员可自行在 Public Desktop 管理。",
    "系统定位、通知、媒体元数据依赖 Windows 权限与具体播放器能力。",
    "天气需要联网；其他核心整理数据保持本地。",
    "便携构建未包含商业代码签名证书和自动更新通道。",
])
add_paragraph(doc, "—— 文档结束 ——")

doc.save(OUTPUT)
print(OUTPUT)
